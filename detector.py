import re
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader


class TemplateDetector:

    def __init__(self, master_file):
        self.master = pd.read_excel(master_file)

    # ----------------------------
    # Read Word Document
    # ----------------------------
    def read_docx(self, file_path):

        doc = Document(file_path)

        text = ""

        # Read paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        # Read footer
        for section in doc.sections:
            footer = section.footer

            for para in footer.paragraphs:
                text += para.text + "\n"

        return text

    # ----------------------------
    # Read PDF
    # ----------------------------
    def read_pdf(self, file_path):

        reader = PdfReader(file_path)

        text = ""

        for page in reader.pages:
            extracted = page.extract_text()

            if extracted:
                text += extracted + "\n"

        return text

    # ----------------------------
    # Extract Template ID
    # ----------------------------
    def extract_template_id(self, text):

        pattern = r"V-QMS-\d+"

        match = re.search(pattern, text, re.IGNORECASE)

        if match:
            return match.group().upper()

        return None

    # ----------------------------
    # Extract Version
    # ----------------------------
    # def extract_version(self, text):

    #     pattern = r"Version\s*:?\s*(\d+\.\d+)"

    #     match = re.search(pattern, text, re.IGNORECASE)

    #     if match:
    #         return match.group(1)

    #     return None
    def extract_version(self, text):

        pattern = r"Version\s*:?\s*(\d+\.\d+)"

        matches = re.findall(pattern, text, re.IGNORECASE)

        # print("All versions found:", matches)

        if matches:
            return matches[-1]   # Last occurrence

        return None
    # ----------------------------
    # Check Master Excel
    # ----------------------------
    def check_template(self, template_id, document_version):

        row = self.master[
            self.master["TemplateID"] == template_id
        ]

        if row.empty:
            return {
                "status": "Not Found",
                "message": "Template ID not found in master list."
            }

        row = row.iloc[0]

        latest_version = float(row["Version"])
        document_version = float(document_version)
        excel_status = row["Status"]

        # If template itself is obsolete/withdrawn
        if excel_status.lower() in ["withdrawn", "superseded", "obsolete"]:

            return {
                "template_id": template_id,
                "document_version": document_version,
                "latest_version": latest_version,
                "status": excel_status,
                "message": f"Template has been {excel_status}."
            }

        # Compare versions
        if document_version == latest_version:

            status = "Latest"

            message = "Template is up-to-date."

        else:

            status = "Outdated"

            message = (
                f"Template is outdated. "
                f"Latest version is {latest_version}."
            )

        return {
        "template_id": template_id,
        "document_version": str(document_version),
        "latest_version": str(latest_version),
        "status": status,
        "message": message
        }

    # ----------------------------
    # Detect Automatically
    # ----------------------------
    def detect(self, file_path):

        if file_path.lower().endswith(".docx"):

            text = self.read_docx(file_path)
            # print(text)

        elif file_path.lower().endswith(".pdf"):

            text = self.read_pdf(file_path)

        else:

            return {
                "error": "Unsupported file format."
            }

        template_id = self.extract_template_id(text)

        if not template_id:
            return {
                "error": "Template ID not found."
            }

        version = self.extract_version(text)

        if not version:
            return {
                "error": "Version not found."
            }

        result = self.check_template(template_id, version)

        return result


# ----------------------------
# Main Program
# ----------------------------

if __name__ == "__main__":

    detector = TemplateDetector("master.xlsx")

    result = detector.detect("sample.docx")

    print(result)